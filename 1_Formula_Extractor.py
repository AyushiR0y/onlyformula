import os
import re
import json
import tempfile
from typing import List, Dict, Tuple, Optional, Set
import streamlit as st
from pdfminer.high_level import extract_text as extract_text_from_pdf_lib
from dotenv import load_dotenv
from dataclasses import dataclass, asdict
from openai import OpenAI
from openai import AzureOpenAI
import hashlib
import time
import traceback
from pathlib import Path
import pandas as pd
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import numpy as np
from collections import defaultdict


load_dotenv()


# Configuration
ALLOWED_EXTENSIONS = {'pdf', 'txt', 'docx'}
MAX_FILE_SIZE = 16 * 1024 * 1024  # 16MB
MAX_VARIANTS = 5


AZURE_API_KEY = os.getenv('AZURE_OPENAI_API_KEY')
AZURE_ENDPOINT = os.getenv('AZURE_OPENAI_ENDPOINT')
DEPLOYMENT_NAME = os.getenv('AZURE_OPENAI_DEPLOYMENT_NAME')
AZURE_API_VERSION = os.getenv('AZURE_OPENAI_API_VERSION')

if not AZURE_API_KEY:
    st.warning("⚠️ **AZURE_OPENAI_API_KEY environment variable not set!** Please set your API key to enable advanced formula extraction from documents.")
    MOCK_MODE = True
    client = None
else:
    try:
        client = AzureOpenAI(
            azure_endpoint=AZURE_ENDPOINT,
            api_key=AZURE_API_KEY,
            api_version=AZURE_API_VERSION
        )
        MOCK_MODE = False

    except Exception as e:
        st.error(f"❌ Failed to initialize Azure OpenAI client: {str(e)}")
        MOCK_MODE = True
        client = None


# --- STABLE CHUNKING CONFIGURATION ---
STABLE_CHUNK_CONFIG = {
    'chunk_size': 1500,
    'chunk_overlap': 300,
    'max_chunks_per_formula': 5,
    'relevance_threshold': 0.3,
    'min_chunk_size': 500,
    'max_context_length': 4000,
}

# --- QUOTA MANAGEMENT CONFIGURATION ---
QUOTA_CONFIG = {
    'max_retries': 3,
    'retry_delay': 5,
    'fallback_enabled': True,
    'batch_size': 5,
    'emergency_stop_after_failures': 10,
}

# --- INPUT VARIABLES DEFINITIONS ---
INPUT_VARIABLES = {
    'TERM_START_DATE': 'Date when the policy starts',
    'FUP_Date': 'First Unpaid Premium date',
    'ENTRY_AGE': 'Age of the policyholder at policy inception',
    'FULL_TERM_PREMIUM': 'Annual Premium amount',
    'BOOKING_FREQUENCY': 'Frequency of premium payments (monthly, quarterly, yearly), used in total premium paid calculation',
    'PREMIUM_TERM': 'Premium Payment Term',
    'SUM_ASSURED': 'Sum Assured - guaranteed amount on maturity/death',
    'Income_Benefit_Amount': 'Amount of income benefit',
    'Income_Benefit_Frequency': 'Frequency of income benefit payout',
    'DATE_OF_SURRENDER': 'Date when policy is surrendered',
    'no_of_premium_paid': 'Years passed since date of commencement till FUP',
    'maturity_date': 'Date of commencement + (BENEFIT_TERM * 12 months)',
    'BENEFIT_TERM': 'The duration (in years) for which the policy benefits are payable',
    'GSV_FACTOR': 'Guaranteed Surrender Value Factor, a percentage used to calculate the minimum guaranteed surrender value from total premiums paid.',
    'SSV1_FACTOR': 'Surrender Value Factor used to compute Special Surrender Value (SSV) related to sum assured on death',
    'SSV3_FACTOR': 'A special factor used to compute Special Surrender Value (SSV) related to paid-up income benefits',
    'SSV2_FACTOR':'A special factor used to compute Special Surrender Value (SSV) related to return of premium (ROP)',
    'FUND_VALUE': 'The total value of the policy fund at the time of surrender or maturity',
    'N':'min(Policy_term, 20) - Elapsed_policy_duration',
    'SYSTEM_PAID': 'The amount paid by the system for surrender or maturity',
    'FUND_FACTOR': 'A factor used in the computation of Surrender Charge based on fund value, Capital fund value and Surrender charge value',
}

BASIC_DERIVED_FORMULAS = {
    'no_of_premium_paid': 'Calculate based on difference between TERM_START_DATE and FUP_Date',
    'policy_year': 'Calculate based on difference between TERM_START_DATE and DATE_OF_SURRENDER + 1',
    'maturity_date': 'TERM_START_DATE + (BENEFIT_TERM* 12) months',
    'Final_surrender_value':'Final surrender value paid',
    'Elapsed_policy_duration': 'How many years have passed since policy start',
    'CAPITAL_FUND_VALUE': 'The total value of the policy fund at the time of surrender or maturity, including any bonuses or additional benefits',
    'FUND_FACTOR': 'A factor used to compute the fund value based on the total premiums paid and the policy term'
}

DEFAULT_TARGET_OUTPUT_VARIABLES = [
    'TOTAL_PREMIUM_PAID',
    'TEN_TIMES_AP',
    'one_oh_five_percent_total_premium',
    'SUM_ASSURED_ON_DEATH',
    'GSV',
    'PAID_UP_SA',
    'PAID_UP_SA_ON_DEATH',
    'PAID_UP_INCOME_INSTALLMENT',
    'SSV1',
    'SSV2',
    'SSV3',
    'SSV',
    'SURRENDER_PAID_AMOUNT',
]

@dataclass
class ExtractedFormula:
    formula_name: str
    formula_expression: str
    variants_info: str
    business_context: str
    source_method: str
    document_evidence: str
    specific_variables: Dict[str, str]
    is_conditional: bool = False
    conditions: List[Dict] = None  # List of {condition: str, expression: str}

    def to_dict(self):
        return asdict(self)

@dataclass
class DocumentExtractionResult:
    input_variables: Dict[str, str]
    basic_derived_formulas: Dict[str, str]
    extracted_formulas: List[ExtractedFormula]

    def to_dict(self):
        return asdict(self)


def get_file_hash(file_bytes: bytes) -> str:
    return hashlib.md5(file_bytes).hexdigest()


@st.cache_data(ttl=3600, show_spinner=False)
def extract_formulas_cached(
    file_hash: str,
    text: str,
    target_outputs: List[str],
    chunk_size: int,
    chunk_overlap: int,
    input_variables: Dict[str, str]
) -> Dict:
    extractor = StableChunkedDocumentFormulaExtractor(
        target_outputs=target_outputs,
        input_variables=input_variables
    )
    result = extractor.extract_formulas_from_document(text)
    return {
        'input_variables': result.input_variables,
        'basic_derived_formulas': result.basic_derived_formulas,
        'extracted_formulas': [
            {
                'formula_name': f.formula_name,
                'formula_expression': f.formula_expression,
                'variants_info': f.variants_info,
                'business_context': f.business_context,
                'source_method': f.source_method,
                'document_evidence': f.document_evidence,
                'specific_variables': f.specific_variables,
                'is_conditional': f.is_conditional,
                'conditions': f.conditions,
            }
            for f in result.extracted_formulas
        ],
    }


def normalize_extracted_formulas(formulas: List[ExtractedFormula]) -> List[ExtractedFormula]:
    normalized = []
    for formula in formulas:
        name_upper = (formula.formula_name or "").upper()
        expression = formula.formula_expression or ""

        if name_upper == "TOTAL_PREMIUM_PAID":
            expression = "FULL_TERM_PREMIUM * no_of_premium_paid * BOOKING_FREQUENCY"

        expression = re.sub(
            r'present_value_of_paid_up_sum_assured_on_death',
            'PAID_UP_SA_ON_DEATH',
            expression,
            flags=re.IGNORECASE
        )

        if expression != formula.formula_expression:
            formula = ExtractedFormula(
                formula_name=formula.formula_name,
                formula_expression=expression,
                variants_info=formula.variants_info,
                business_context=formula.business_context,
                source_method=formula.source_method,
                document_evidence=formula.document_evidence,
                specific_variables=formula.specific_variables,
                is_conditional=formula.is_conditional,
                conditions=formula.conditions,
            )
        normalized.append(formula)
    return normalized


def compare_formula_expressions(expr1: str, expr2: str) -> bool:
    """Returns True if expressions differ (ignoring whitespace/case)"""
    def normalize(e):
        return re.sub(r'\s+', '', (e or "").strip().lower())
    return normalize(expr1) != normalize(expr2)


def build_formula_comparison_table(all_variant_formulas: Dict[str, List[Dict]], variant_names: List[str]) -> pd.DataFrame:
    """Build a comparison dataframe across variants."""
    all_formula_names = []
    for variant_name, formulas in all_variant_formulas.items():
        for f in formulas:
            name = f.get("formula_name", "")
            if name and name not in all_formula_names:
                all_formula_names.append(name)

    rows = []
    for formula_name in all_formula_names:
        row = {"Formula Name": formula_name}
        expressions = {}
        for variant_name in variant_names:
            formulas = all_variant_formulas.get(variant_name, [])
            match = next((f for f in formulas if f.get("formula_name", "") == formula_name), None)
            expr = match["formula_expression"] if match else "—"
            row[variant_name] = expr
            expressions[variant_name] = expr

        # Mark if any differ
        unique_exprs = set(e for e in expressions.values() if e != "—")
        row["_has_diff"] = len(unique_exprs) > 1
        rows.append(row)

    return pd.DataFrame(rows)


class StableChunkedDocumentFormulaExtractor:
    """Extracts formulas from large documents using stable chunking ratios"""

    def __init__(self, target_outputs: List[str], input_variables: Dict[str, str] = None):
        self.input_variables = input_variables if input_variables is not None else INPUT_VARIABLES
        self.basic_derived = BASIC_DERIVED_FORMULAS
        self.target_outputs = target_outputs
        self.config = STABLE_CHUNK_CONFIG
        self.quota_config = QUOTA_CONFIG
        self.failure_count = 0

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config['chunk_size'],
            chunk_overlap=self.config['chunk_overlap'],
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", " ", ""]
        )

        self.formula_keywords = {
            'high_priority': ['surrender', 'gsv', 'ssv', 'formula', 'calculate', 'premium', 'benefit'],
            'medium_priority': ['paid-up', 'maturity', 'death', 'sum assured', 'charge', 'value'],
            'low_priority': ['policy', 'term', 'amount', 'date', 'factor', 'rate']
        }

    def _extract_formula_offline(self, formula_name: str, context: str) -> Optional[ExtractedFormula]:
        formula_patterns = {
            'equals': r'([A-Z_]+)\s*[=:]\s*([^.\n]+)',
            'calculation': r'([A-Z_]+)\s*(?:is calculated|calculated as|=|:)\s*([^.\n]+)',
            'formula': r'(?:formula|calculation)\s*(?:for|of)\s*([A-Z_]+)[:\s]*([^.\n]+)',
            'definition': r'([A-Z_]+)\s*(?:means|refers to|defined as)\s*([^.\n]+)'
        }
        formula_lower = formula_name.lower()
        context_lower = context.lower()
        if formula_lower in context_lower:
            sentences = context.split('.')
            relevant_sentences = [s.strip() for s in sentences if formula_lower in s.lower()]
            if relevant_sentences:
                for sentence in relevant_sentences:
                    for pattern_name, pattern in formula_patterns.items():
                        matches = re.findall(pattern, sentence, re.IGNORECASE)
                        if matches:
                            formula_expr = matches[0][1] if len(matches[0]) > 1 else matches[0][0]
                            variables_found = [v for v in self.input_variables.keys() if v.lower() in formula_expr.lower()]
                            return ExtractedFormula(
                                formula_name=formula_name.upper(),
                                formula_expression=formula_expr.strip(),
                                variants_info="Extracted using offline pattern matching",
                                business_context=f"Offline extraction for {formula_name}",
                                source_method='offline_pattern_matching',
                                document_evidence=sentence[:200] + "..." if len(sentence) > 200 else sentence,
                                specific_variables={var: self.input_variables[var] for var in variables_found},
                                is_conditional=False,
                                conditions=None,
                            )
        return None

    def extract_formulas_from_document(self, text: str) -> DocumentExtractionResult:
        if MOCK_MODE or not AZURE_API_KEY:
            return self._explain_no_extraction()

        try:
            progress_bar = st.progress(0)
            progress_bar.progress(10)
            chunks = self._create_stable_chunks(text)
            progress_bar.progress(20)
            scored_chunks = self._score_chunks_for_relevance(chunks)

            if not self._check_api_status():
                st.error("❌ API is not accessible. Using offline extraction only.")
                return self._fallback_to_offline_extraction(text)

            progress_bar.progress(30)
            extracted_formulas = []
            total_formulas = len(self.target_outputs)

            batch_size = self.quota_config['batch_size']
            for batch_start in range(0, len(self.target_outputs), batch_size):
                batch_end = min(batch_start + batch_size, len(self.target_outputs))
                batch_formulas = self.target_outputs[batch_start:batch_end]

                for i, formula_name in enumerate(batch_formulas):
                    overall_progress = 30 + int(((batch_start + i) / total_formulas) * 60)
                    progress_bar.progress(overall_progress)

                    if self.failure_count >= self.quota_config['emergency_stop_after_failures']:
                        st.error(f"🛑 Emergency stop: {self.failure_count} consecutive failures.")
                        remaining_formulas = self.target_outputs[batch_start + i:]
                        for remaining_formula in remaining_formulas:
                            offline_result = self._extract_formula_offline(remaining_formula, text)
                            if offline_result:
                                extracted_formulas.append(offline_result)
                        break

                    formula_result = self._extract_formula_stable(formula_name, scored_chunks, text)
                    if formula_result:
                        extracted_formulas.append(formula_result)
                        self.failure_count = 0
                    else:
                        self.failure_count += 1
                    time.sleep(0.5 if formula_result else 1.0)

                if self.failure_count >= self.quota_config['emergency_stop_after_failures']:
                    break
                if batch_end < len(self.target_outputs):
                    time.sleep(self.quota_config['retry_delay'])

            progress_bar.progress(100)
            return DocumentExtractionResult(
                input_variables=self.input_variables,
                basic_derived_formulas=self.basic_derived,
                extracted_formulas=extracted_formulas,
            )

        except Exception as e:
            st.error(f"❌ **Stable chunked extraction failed:** {e}")
            st.exception(e)
            return self._explain_no_extraction()

    def _check_api_status(self) -> bool:
        try:
            client.chat.completions.create(
                model=DEPLOYMENT_NAME,
                messages=[{"role": "user", "content": "Test"}],
            )
            return True
        except Exception as e:
            st.warning(f"API check failed: {e}")
            return False

    def _fallback_to_offline_extraction(self, text: str) -> DocumentExtractionResult:
        st.info("🔄 Falling back to offline pattern-based extraction...")
        extracted_formulas = []
        for formula_name in self.target_outputs:
            offline_result = self._extract_formula_offline(formula_name, text)
            if offline_result:
                extracted_formulas.append(offline_result)
        return DocumentExtractionResult(
            input_variables=self.input_variables,
            basic_derived_formulas=self.basic_derived,
            extracted_formulas=extracted_formulas,
        )

    def _create_stable_chunks(self, text: str) -> List[Dict]:
        text_chunks = self.text_splitter.split_text(text)
        stable_chunks = []
        for i, chunk_text in enumerate(text_chunks):
            if len(chunk_text) >= self.config['min_chunk_size']:
                stable_chunks.append({
                    'id': i,
                    'text': chunk_text,
                    'char_count': len(chunk_text),
                    'word_count': len(chunk_text.split()),
                    'relevance_score': 0.0,
                    'chunk_ratio': len(chunk_text) / len(text),
                    'position_ratio': i / len(text_chunks),
                })
        return stable_chunks

    def _score_chunks_for_relevance(self, chunks: List[Dict]) -> List[Dict]:
        for chunk in chunks:
            text_lower = chunk['text'].lower()
            score = 0.0
            for keyword in self.formula_keywords['high_priority']:
                score += text_lower.count(keyword) * 3
            for keyword in self.formula_keywords['medium_priority']:
                score += text_lower.count(keyword) * 2
            for keyword in self.formula_keywords['low_priority']:
                score += text_lower.count(keyword) * 1
            normalized_score = score / (chunk['word_count'] / 100)
            position_bonus = 1.0 - (chunk['position_ratio'] * 0.3)
            chunk['relevance_score'] = normalized_score * position_bonus
        return sorted(chunks, key=lambda x: x['relevance_score'], reverse=True)

    def _extract_formula_stable(self, formula_name: str, scored_chunks: List[Dict], full_text: str) -> Optional[ExtractedFormula]:
        relevant_chunks = self._select_relevant_chunks_for_formula(formula_name, scored_chunks)
        if not relevant_chunks:
            return None
        combined_context = self._combine_chunks_stable(relevant_chunks)
        return self._extract_formula_with_context(formula_name, combined_context)

    def _select_relevant_chunks_for_formula(self, formula_name: str, scored_chunks: List[Dict]) -> List[Dict]:
        formula_specific_keywords = {
            'surrender': ['surrender', 'gsv', 'ssv', 'cash', 'quit', 'capital units', 'elapsed', 'policy term', '1.05', 'three years', 'redemption'],
            'surrender_charge': ['surrender charge', 'capital units', '1.05', 'elapsed policy duration', 'policy term', 'three years', 'redemption'],
            'premium': ['premium', 'payment', 'annual', 'monthly'],
            'benefit': ['benefit', 'payout', 'income', 'amount'],
            'death': ['death', 'mortality', 'sum assured'],
            'maturity': ['maturity', 'endowment', 'maturity date'],
            'charge': ['charge', 'fee', 'deduction', 'cost']
        }
        relevant_keywords = []
        for category, keywords in formula_specific_keywords.items():
            if any(keyword in formula_name.lower() for keyword in keywords):
                relevant_keywords.extend(keywords)

        formula_scored_chunks = []
        for chunk in scored_chunks:
            text_lower = chunk['text'].lower()
            formula_score = chunk['relevance_score']
            for keyword in relevant_keywords:
                if keyword in text_lower:
                    formula_score += 2.0
            if formula_name.lower() in text_lower:
                formula_score += 5.0
            chunk_copy = chunk.copy()
            chunk_copy['formula_score'] = formula_score
            formula_scored_chunks.append(chunk_copy)

        formula_scored_chunks.sort(key=lambda x: x['formula_score'], reverse=True)
        selected_chunks = []
        for chunk in formula_scored_chunks:
            if (len(selected_chunks) < self.config['max_chunks_per_formula'] and
                    chunk['formula_score'] >= self.config['relevance_threshold']):
                selected_chunks.append(chunk)
        return selected_chunks

    def _combine_chunks_stable(self, chunks: List[Dict]) -> str:
        combined_text = ""
        current_length = 0
        max_length = self.config['max_context_length']
        for chunk in chunks:
            chunk_text = chunk['text']
            if current_length + len(chunk_text) > max_length:
                remaining_space = max_length - current_length
                if remaining_space > 200:
                    chunk_text = chunk_text[:remaining_space] + "..."
                    combined_text += f"\n\n--- Chunk {chunk['id']} ---\n{chunk_text}"
                break
            combined_text += f"\n\n--- Chunk {chunk['id']} ---\n{chunk_text}"
            current_length += len(chunk_text)
        return combined_text

    def _extract_formula_with_context(self, formula_name: str, context: str) -> Optional[ExtractedFormula]:
        prompt = f"""
Extract the calculation formula for "{formula_name}" from the following document content.

DOCUMENT CONTENT:
{context}

AVAILABLE VARIABLES:
{', '.join(self.input_variables.keys())}

INSTRUCTIONS:
1. Identify a mathematical formula or calculation method for "{formula_name}"
2. Use the available variables and generated variables from the previous formulas where possible
3. Extract the formula expression from natural language
4. Look carefully at variable name suffixes like "_ON_DEATH" - use the correct variant
5. IMPORTANT: You MUST provide a formula for "{formula_name}". Do not skip it.
6. If the exact formula is not clearly defined in the document:
   - Make a reasonable inference based on similar formulas
   - Use industry-standard calculations as fallback
   - Provide a placeholder formula with low confidence
7. Pay close attention to formulas involving:
   - Terms around GSV, SSV (Surrender Paid Amount is usually a max of multiple components)
   - Exponential terms like (1/1.05)^N
   - Conditions like policy term > 3 years
   - Capital Units references
   - ON_DEATH is an important qualifier
8. Reuse PAID_UP_SA_ON_DEATH in future formulas instead of Present_Value_of_paid_up_sum_assured_on_death
9. For PAID_UP_INCOME_INSTALLMENT: Income_Benefit_Amount * Income_Benefit_Frequency is always used, along with no_of_premium_paid and PREMIUM_TERM
10. Total Premium paid uses FULL_TERM_PREMIUM, no_of_premium_paid and BOOKING_FREQUENCY

TABLE EXTRACTION GUIDELINES (CRITICAL FOR STRUCTURED DATA):
11. **Tables and Merged Cells**: If the document contains tables:
    - Look for MERGED ROWS or MERGED CELLS that act as category headers or group titles
    - A merged cell/row typically spans multiple columns and contains a formula name or category
    - Formulas may be listed in rows BELOW the merged header, with conditions in different columns
    - Example table structure:
      | Formula Name (merged across columns) | Condition 1 | Condition 2 | Condition 3 |
      | Specific Case A                       | expr_A1     | expr_A2     | expr_A3     |
      | Specific Case B                       | expr_B1     | expr_B2     | expr_B3     |
    
12. **Column Headers as Conditions**: 
    - Column headers often represent CONDITIONS (e.g., "Year < 3", "Year 3-5", "Year > 5")
    - If the target formula appears as a merged row header, extract formulas from cells below it
    - Map each column header to its corresponding condition
    - Combine row labels (left column) with column conditions if both exist
    
13. **Row Labels and Grouping**:
    - First column often contains row labels that qualify the formula further
    - Combine row labels with the merged title to understand the complete context
    - Example: Merged title "GSV", Row label "If premium paid < 3 years" → GSV calculation for that condition
    
14. **Multiple Formula Variants in Tables**:
    - A single formula name may have multiple variants based on table structure
    - Look for patterns like: "Formula X (Case 1)", "Formula X (Case 2)" in rows
    - Or conditions specified in column headers applying to different expressions
    - Extract ALL variants and combine them using appropriate conditional logic
    
15. **Table Context Awareness**:
    - Read table captions, notes, and legends - they often contain critical multipliers or factors
    - Footer notes may specify default values, exceptions, or global conditions
    - Adjacent cells may contain explanatory text that clarifies the formula
    
16. **Handling Table Fragments**:
    - Tables may be split across chunks - look for continuation patterns
    - If you see partial table structure, infer the pattern from visible rows
    - Headers like "continued..." or "...cont'd" indicate table continuation

MULTI-LEVEL & CONDITIONAL FORMULA HANDLING (CRITICAL):
- If the formula has IF/ELSE conditions (e.g., "if policy_year < 3 then X else Y"), extract ALL branches
- If the formula uses a tiered/stepped structure (e.g., different rates for different durations), list each tier
- If the formula is a piecewise function, capture every piece with its condition
- Capture nested conditions (e.g., "if A then (if B then X else Y) else Z")
- For MAX/MIN selections, include all candidates
- Use Python-style conditional syntax: "X if condition else Y"
- For multi-step formulas, show each intermediate step clearly labelled

Examples:
- "surrender value is higher of GSV or SSV" → "MAX(GSV, SSV)"
- "if policy_year < 3: 0; elif policy_year < 5: GSV_FACTOR * TOTAL_PREMIUM_PAID * 0.5; else: GSV_FACTOR * TOTAL_PREMIUM_PAID"
- "Sum Assured on Death is higher of SA, 10x AP or ROP" → "MAX(SUM_ASSURED, TEN_TIMES_AP, one_oh_five_percent_total_premium)"
- Conditional: "PAID_UP_SA * SSV1_FACTOR if no_of_premium_paid >= 3 else 0"

TABLE-BASED FORMULA EXAMPLES:
- Table with merged header "GSV" and rows showing different policy years:
  * Row 1: "Year < 3" | "0"
  * Row 2: "Year 3-5" | "0.5 * TOTAL_PREMIUM_PAID * GSV_FACTOR"
  * Row 3: "Year > 5" | "TOTAL_PREMIUM_PAID * GSV_FACTOR"
  → Extract as: "0 if policy_year < 3 else (0.5 * TOTAL_PREMIUM_PAID * GSV_FACTOR if policy_year <= 5 else TOTAL_PREMIUM_PAID * GSV_FACTOR)"

- Table with merged title "SURRENDER_VALUE" spanning columns with different product types:
  * Column headers: "Product A" | "Product B" | "Product C"
  * If searching for SURRENDER_VALUE and context matches Product A → extract from that column
  
- Table showing formulas with row groupings:
  * Merged row: "Special Surrender Value (SSV)"
  * Sub-row 1: "Component 1" | "PAID_UP_SA * SSV1_FACTOR"
  * Sub-row 2: "Component 2" | "ROP * SSV2_FACTOR"
  * Sub-row 3: "Component 3" | "PAID_UP_INCOME * SSV3_FACTOR"
  * If searching for SSV → "SSV1 + SSV2 + SSV3" or "MAX(SSV1, SSV2, SSV3)" depending on table context

RESPONSE FORMAT:
FORMULA_EXPRESSION: [primary mathematical expression; if conditional, use Python-style if/else inline or multi-line]
IS_CONDITIONAL: [YES or NO]
CONDITIONS:
- CONDITION_1: [condition text e.g. "policy_year >= 3"] | EXPRESSION_1: [formula when condition met]
- CONDITION_2: [condition text e.g. "policy_year < 3"] | EXPRESSION_2: [formula when condition met]
(Only fill CONDITIONS block if IS_CONDITIONAL is YES, otherwise leave blank)
VARIABLES_USED: [comma-separated list]
DOCUMENT_EVIDENCE: [exact or near-exact text passage from the document that supports this formula — quote it verbatim if possible. If from a table, include: table title/caption, merged row header, column headers, and relevant cell values. Write "INFERRED" only if not explicit in document. Example: "From Table 3.2 'Surrender Values': Row header 'GSV' with columns 'Year<3: 0', 'Year 3-5: 0.5*Premium*Factor', 'Year>5: Premium*Factor'"]
BUSINESS_CONTEXT: [brief explanation of what this formula calculates]

Respond with only the requested format.
"""

        models_to_try = ["gpt-3.5-turbo", "gpt-4o-mini", "gpt-4"]

        for model in models_to_try:
            try:
                response = client.chat.completions.create(
                    model=DEPLOYMENT_NAME,
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=1200,
                    temperature=0.1,
                    top_p=0.95
                )
                response_text = response.choices[0].message.content
                parsed_formula = self._parse_stable_formula_response(response_text, formula_name)
                if parsed_formula:
                    return parsed_formula

            except Exception as e:
                error_msg = str(e)
                if "429" in error_msg or "quota" in error_msg.lower():
                    time.sleep(2)
                    continue
                elif "404" in error_msg or "model" in error_msg.lower():
                    continue
                else:
                    st.error(f"❌ Error with {model} for {formula_name}: {e}")
                    continue

        offline_result = self._extract_formula_offline(formula_name, context)
        if offline_result is None:
            return self._create_placeholder_formula(formula_name, "No extraction possible")
        return offline_result

    def _create_placeholder_formula(self, formula_name: str, reason: str) -> ExtractedFormula:
        return ExtractedFormula(
            formula_name=formula_name.upper(),
            formula_expression=f"{formula_name}  # {reason}",
            variants_info="Placeholder - could not extract",
            business_context=f"Placeholder for {formula_name}",
            source_method='placeholder',
            document_evidence="No evidence found",
            specific_variables={},
            is_conditional=False,
            conditions=None,
        )

    def _parse_stable_formula_response(self, response_text: str, formula_name: str) -> Optional[ExtractedFormula]:
        try:
            formula_match = re.search(r'FORMULA_EXPRESSION:\s*(.+?)(?=\nIS_CONDITIONAL|\nVARIABLES_USED|$)', response_text, re.DOTALL | re.IGNORECASE)
            formula_expression = formula_match.group(1).strip() if formula_match else "Formula not clearly defined"

            is_conditional_match = re.search(r'IS_CONDITIONAL:\s*(YES|NO)', response_text, re.IGNORECASE)
            is_conditional = is_conditional_match and is_conditional_match.group(1).upper() == "YES"

            conditions = None
            if is_conditional:
                conditions_block = re.search(r'CONDITIONS:\s*(.+?)(?=\nVARIABLES_USED|$)', response_text, re.DOTALL | re.IGNORECASE)
                if conditions_block:
                    conditions_text = conditions_block.group(1)
                    conditions = []
                    cond_lines = re.findall(r'-\s*CONDITION_\d+:\s*(.+?)\s*\|\s*EXPRESSION_\d+:\s*(.+)', conditions_text)
                    for cond, expr in cond_lines:
                        conditions.append({"condition": cond.strip(), "expression": expr.strip()})

            variables_match = re.search(r'VARIABLES_USED:\s*(.+?)(?=\nDOCUMENT_EVIDENCE|$)', response_text, re.DOTALL | re.IGNORECASE)
            variables_str = variables_match.group(1).strip() if variables_match else ""
            specific_variables = self._parse_variables_stable(variables_str)

            evidence_match = re.search(r'DOCUMENT_EVIDENCE:\s*(.+?)(?=\nBUSINESS_CONTEXT|$)', response_text, re.DOTALL | re.IGNORECASE)
            document_evidence = evidence_match.group(1).strip() if evidence_match else "No supporting evidence found"

            context_match = re.search(r'BUSINESS_CONTEXT:\s*(.+?)$', response_text, re.DOTALL | re.IGNORECASE)
            business_context = context_match.group(1).strip() if context_match else f"Calculation for {formula_name}"

            return ExtractedFormula(
                formula_name=formula_name.upper(),
                formula_expression=formula_expression,
                variants_info="Extracted using stable chunking approach",
                business_context=business_context,
                source_method='stable_chunked_extraction',
                document_evidence=document_evidence[:800],
                specific_variables=specific_variables,
                is_conditional=is_conditional,
                conditions=conditions,
            )

        except Exception as e:
            st.error(f"Error parsing response for {formula_name}: {e}")
            return None

    def _parse_variables_stable(self, variables_str: str) -> Dict[str, str]:
        specific_variables = {}
        if variables_str:
            var_names = [var.strip().upper() for var in variables_str.split(',')]
            for var_name in var_names:
                if var_name in self.input_variables:
                    specific_variables[var_name] = self.input_variables[var_name]
                elif var_name in self.basic_derived:
                    specific_variables[var_name] = self.basic_derived[var_name]
                else:
                    for input_var in self.input_variables:
                        if var_name in input_var or input_var in var_name:
                            specific_variables[input_var] = self.input_variables[input_var]
                            break
        return specific_variables

    def _explain_no_extraction(self) -> DocumentExtractionResult:
        return DocumentExtractionResult(
            input_variables=self.input_variables,
            basic_derived_formulas=self.basic_derived,
            extracted_formulas=[],
        )


def extract_text_from_file(file_bytes, file_extension):
    try:
        if file_extension == '.pdf':
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
                tmp_file.write(file_bytes)
                tmp_file_path = tmp_file.name
            try:
                # Use layout-aware extraction to better preserve table structure
                from pdfminer.layout import LAParams
                laparams = LAParams(
                    line_margin=0.5,
                    word_margin=0.1,
                    char_margin=2.0,
                    boxes_flow=0.5,
                    detect_vertical=True,
                    all_texts=False
                )
                text = extract_text_from_pdf_lib(tmp_file_path, laparams=laparams)
                os.unlink(tmp_file_path)
                return text
            except Exception as e:
                os.unlink(tmp_file_path)
                raise e
        elif file_extension == '.txt':
            return file_bytes.decode('utf-8')
        elif file_extension == '.docx':
            try:
                import docx
                from io import BytesIO
                doc = docx.Document(BytesIO(file_bytes))
                
                # Extract both paragraphs and tables
                full_text = []
                for element in doc.element.body:
                    if element.tag.endswith('p'):  # Paragraph
                        para = next((p for p in doc.paragraphs if p._element == element), None)
                        if para and para.text.strip():
                            full_text.append(para.text)
                    elif element.tag.endswith('tbl'):  # Table
                        table = next((t for t in doc.tables if t._element == element), None)
                        if table:
                            full_text.append("\n--- TABLE START ---")
                            for i, row in enumerate(table.rows):
                                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                                if i == 0:
                                    full_text.append(f"HEADER: {row_text}")
                                else:
                                    full_text.append(row_text)
                            full_text.append("--- TABLE END ---\n")
                
                return '\n'.join(full_text)
            except ImportError:
                st.error("File type not recognised or required library not installed.")
                return ""
        else:
            return ""
    except Exception as e:
        st.error(f"Error extracting text from file: {e}")
        return ""


def load_css(file_name="style.css"):
    current_dir = os.path.dirname(os.path.abspath(__file__))
    css_path = os.path.join(current_dir, file_name)
    if os.path.exists(css_path):
        with open(css_path, 'r') as f:
            st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)


def render_formula_card(formula: Dict, variant_name: str = None, highlight: bool = False, all_variants_formulas: Dict = None):
    """Render a single formula card with evidence and conditional display."""
    import html
    import hashlib
    
    expr = formula.get("formula_expression", "")
    name = formula.get("formula_name", "")
    evidence = formula.get("document_evidence", "")
    context = formula.get("business_context", "")
    is_conditional = formula.get("is_conditional", False)
    conditions = formula.get("conditions") or []

    # Escape all text that goes into HTML
    safe_name = html.escape(name)
    safe_context = html.escape(context) if context else ""
    
    # Create unique key for this formula card
    variant_prefix = f"{variant_name}_" if variant_name else ""
    unique_key = f"{variant_prefix}{name}_{hashlib.md5(expr.encode()).hexdigest()[:8]}"
    
    border_color = "#e74c3c" if highlight else "#3498db"
    badge = "<span style='background: #e74c3c; color: white; padding: 2px 8px; border-radius: 10px; font-size: 0.7rem; font-weight: 600; margin-left: 8px;'>⚠ DIFFERS</span>" if highlight else ""
    
    # Compact header - single line with badge
    st.markdown(f"<div style='border-left: 3px solid {border_color}; padding: 8px 12px; margin: 8px 0; background: #fafafa;'><strong style='color: #2c3e50; font-size: 0.95rem;'>{safe_name}</strong>{badge}</div>", unsafe_allow_html=True)
    
    # Formula expression - compact
    st.code(expr, language="python")

    # Compact additional info
    extras = []
    if is_conditional and conditions:
        extras.append("📊 Conditional")
    if evidence and evidence not in ("No supporting evidence found", "No evidence found", "INFERRED"):
        extras.append("📄 Has Evidence")
    if context and context != f"Calculation for {name}":
        extras.append(f"💡 {safe_context[:80]}..." if len(safe_context) > 80 else f"💡 {safe_context}")
    
    if extras:
        st.caption(" • ".join(extras))
    
    # Expandable details only if needed
    details_needed = (is_conditional and conditions) or (evidence and evidence not in ("No supporting evidence found", "No evidence found", "INFERRED"))
    
    if details_needed:
        with st.expander("Show Details", expanded=False):
            if is_conditional and conditions:
                st.markdown("**Conditional Branches:**")
                for idx, cond in enumerate(conditions):
                    st.text(f"• {cond.get('condition','')}")
                    st.code(cond.get('expression', ''), language="python")
            
            if evidence and evidence not in ("No supporting evidence found", "No evidence found", "INFERRED"):
                st.markdown("**Document Evidence:**")
                st.text_area("Source", value=evidence, height=100, disabled=True, label_visibility="collapsed", key=f"evidence_{unique_key}")
    
    st.markdown('<div style="height: 4px;"></div>', unsafe_allow_html=True)


def run_extraction_for_variant(variant_name: str, uploaded_files: list, target_outputs: List[str], input_variables: Dict[str, str]) -> Optional[List[Dict]]:
    """Run extraction for a single variant from one or more uploaded files."""
    if not uploaded_files:
        return None

    combined_text = ""
    for uf in uploaded_files:
        file_bytes = uf.read()
        file_extension = Path(uf.name).suffix.lower()
        text = extract_text_from_file(file_bytes, file_extension)
        combined_text += "\n\n" + text

    if not combined_text.strip():
        st.error(f"Could not extract text for variant '{variant_name}'")
        return None

    if not MOCK_MODE and AZURE_API_KEY:
        file_hash = hashlib.md5(combined_text.encode()).hexdigest()
        target_key = "_".join(sorted(target_outputs))
        cache_key = f"{file_hash}_{target_key}"
        try:
            result_dict = extract_formulas_cached(
                file_hash=cache_key,
                text=combined_text,
                target_outputs=target_outputs,
                chunk_size=STABLE_CHUNK_CONFIG['chunk_size'],
                chunk_overlap=STABLE_CHUNK_CONFIG['chunk_overlap'],
                input_variables=input_variables
            )
            extracted_formulas = [
                ExtractedFormula(
                    formula_name=f['formula_name'],
                    formula_expression=f['formula_expression'],
                    variants_info=f['variants_info'],
                    business_context=f['business_context'],
                    source_method=f['source_method'],
                    document_evidence=f['document_evidence'],
                    specific_variables=f['specific_variables'],
                    is_conditional=f.get('is_conditional', False),
                    conditions=f.get('conditions'),
                )
                for f in result_dict['extracted_formulas']
            ]
            extracted_formulas = normalize_extracted_formulas(extracted_formulas)
        except Exception:
            extractor = StableChunkedDocumentFormulaExtractor(
                target_outputs=target_outputs,
                input_variables=input_variables
            )
            result = extractor.extract_formulas_from_document(combined_text)
            extracted_formulas = normalize_extracted_formulas(result.extracted_formulas)

        return [
            {
                "formula_name": f.formula_name,
                "formula_expression": f.formula_expression,
                "business_context": f.business_context,
                "document_evidence": f.document_evidence,
                "is_conditional": f.is_conditional,
                "conditions": f.conditions,
                "source_method": f.source_method,
                "specific_variables": f.specific_variables,
            }
            for f in extracted_formulas
        ]
    else:
        st.warning("Cannot perform extraction without a configured AZURE_OPENAI_API_KEY.")
        return None


def main():
    st.set_page_config(
        page_title="Document Formula Extractor",
        page_icon="🧮",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    load_css()

    st.markdown(
        """
        <div class="header-container">
            <div class="header-bar">
                <img src="https://raw.githubusercontent.com/AyushiR0y/streamlit_formulagen/main/assets/logo.png">
                <div class="header-title">Document Formula Extractor</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    # --- Session State Init ---
    for key, default in [
        ('extraction_result', None),
        ('selected_output_variables', DEFAULT_TARGET_OUTPUT_VARIABLES.copy()),
        ('custom_output_variable', ""),
        ('user_defined_output_variables', []),
        ('formulas', []),
        ('formulas_saved', False),
        ('editing_formula', -1),
        ('previous_selected_variables', DEFAULT_TARGET_OUTPUT_VARIABLES.copy()),
        ('variant_results', {}),      # {variant_name: List[Dict]}
        ('num_variants', 2),
        ('variant_names', ['Product A', 'Product B']),  # Initialize with at least 2 variants
        ('mode', 'single'),           # 'single' or 'multi'
        ('custom_input_variables', INPUT_VARIABLES.copy()),  # User-customizable input variables
        ('editing_input_var', None),  # Variable name being edited
    ]:
        if key not in st.session_state:
            st.session_state[key] = default

    # Ensure variant_names always has enough elements for num_variants
    while len(st.session_state.variant_names) < st.session_state.num_variants:
        st.session_state.variant_names.append(f"Variant {len(st.session_state.variant_names) + 1}")

    st.markdown("---")

    # ========== KEYWORD SELECTION ==========
    col1, col2 = st.columns([1, 1])

    with col1:
        st.subheader("Select Keywords")
        st.markdown("Choose the keywords for which formulas will be extracted.")

        all_possible_output_variables = sorted(list(set(DEFAULT_TARGET_OUTPUT_VARIABLES + st.session_state.user_defined_output_variables)))
        current_selection = st.multiselect(
            "Target Keywords:",
            options=all_possible_output_variables,
            default=st.session_state.selected_output_variables,
            help="These are the keywords for which the system will try to find formulas."
        )

        if current_selection != st.session_state.previous_selected_variables:
            st.session_state.selected_output_variables = current_selection
            st.session_state.previous_selected_variables = current_selection.copy()
            st.session_state.extraction_result = None
            st.session_state.formulas = []
            st.session_state.formulas_saved = False
            st.session_state.editing_formula = -1
            st.session_state.variant_results = {}
            st.rerun()
        else:
            st.session_state.selected_output_variables = current_selection

        st.session_state.custom_output_variable = st.text_input(
            "Add a custom keyword:",
            value=st.session_state.custom_output_variable,
            key="custom_output_input"
        )
        if st.button("Add Custom Keyword", key="add_custom_formula_button"):
            new_var = st.session_state.custom_output_variable.strip()
            if new_var and new_var not in st.session_state.user_defined_output_variables and new_var not in DEFAULT_TARGET_OUTPUT_VARIABLES:
                st.session_state.user_defined_output_variables.append(new_var)
                if new_var not in st.session_state.selected_output_variables:
                    st.session_state.selected_output_variables.append(new_var)
                    st.session_state.previous_selected_variables = st.session_state.selected_output_variables.copy()
                st.session_state.custom_output_variable = ""
                st.session_state.extraction_result = None
                st.session_state.formulas = []
                st.session_state.formulas_saved = False
                st.session_state.editing_formula = -1
                st.session_state.variant_results = {}
                st.success(f"'{new_var}' added!")
                st.rerun()
            elif new_var:
                st.info(f"'{new_var}' is already in the list.")

    with col2:
        st.subheader("Reference Variables")
        with st.expander("Input Variables (Policy Parameters)", expanded=False):
            st.markdown("**Manage Input Variables** - Edit or add custom variables below.")
            
            # Display existing input variables with edit/delete options
            for var_name, var_desc in sorted(st.session_state.custom_input_variables.items()):
                var_col1, var_col2, var_col3, var_col4 = st.columns([2, 4, 1, 1])
                
                with var_col1:
                    if st.session_state.editing_input_var == var_name:
                        new_var_name = st.text_input("Name", value=var_name, key=f"edit_name_{var_name}", label_visibility="collapsed")
                    else:
                        st.markdown(f"**{var_name}**")
                
                with var_col2:
                    if st.session_state.editing_input_var == var_name:
                        new_var_desc = st.text_input("Description", value=var_desc, key=f"edit_desc_{var_name}", label_visibility="collapsed")
                    else:
                        st.markdown(f"_{var_desc}_")
                
                with var_col3:
                    if st.session_state.editing_input_var == var_name:
                        if st.button("💾", key=f"save_{var_name}", help="Save changes"):
                            if new_var_name.strip():
                                # Remove old entry and add new one
                                if new_var_name != var_name:
                                    del st.session_state.custom_input_variables[var_name]
                                st.session_state.custom_input_variables[new_var_name.strip()] = new_var_desc.strip()
                                st.session_state.editing_input_var = None
                                st.rerun()
                    else:
                        if st.button("✏️", key=f"edit_btn_{var_name}", help="Edit variable"):
                            st.session_state.editing_input_var = var_name
                            st.rerun()
                
                with var_col4:
                    if st.session_state.editing_input_var == var_name:
                        if st.button("❌", key=f"cancel_{var_name}", help="Cancel"):
                            st.session_state.editing_input_var = None
                            st.rerun()
                    else:
                        if st.button("🗑️", key=f"delete_{var_name}", help="Delete variable"):
                            del st.session_state.custom_input_variables[var_name]
                            st.rerun()
            
            # Add new input variable section
            st.markdown("---")
            st.markdown("**Add New Input Variable**")
            new_iv_col1, new_iv_col2, new_iv_col3 = st.columns([2, 4, 1])
            
            with new_iv_col1:
                new_input_var_name = st.text_input("Variable Name", key="new_input_var_name", placeholder="e.g., POLICY_TERM", label_visibility="collapsed")
            
            with new_iv_col2:
                new_input_var_desc = st.text_input("Description", key="new_input_var_desc", placeholder="Description of the variable", label_visibility="collapsed")
            
            with new_iv_col3:
                if st.button("➕ Add", key="add_input_var_btn"):
                    if new_input_var_name.strip():
                        st.session_state.custom_input_variables[new_input_var_name.strip()] = new_input_var_desc.strip()
                        st.success(f"Added '{new_input_var_name.strip()}'")
                        st.rerun()
                    else:
                        st.warning("Please enter a variable name")
        
        with st.expander("Basic Derived Formulas (Pre-computed)", expanded=False):
            derived_data = [{"Variable": name, "Description": desc} for name, desc in BASIC_DERIVED_FORMULAS.items()]
            st.dataframe(pd.DataFrame(derived_data), use_container_width=True, hide_index=True)
        with st.expander("Currently Selected Target Output Variables", expanded=True):
            current_targets = sorted(list(set(st.session_state.selected_output_variables)))
            if current_targets:
                st.dataframe(pd.DataFrame([{"Target Variable": v} for v in current_targets]), use_container_width=True, hide_index=True)
            else:
                st.info("No target formulas selected yet.")

    st.markdown("---")

    # ========== UPLOAD MODE SELECTOR ==========
    st.subheader("📂 Upload Product Documents")
    
    # Info about table handling
    st.info("📊 **Enhanced Table Support**: This extractor now intelligently handles tables with merged cells, row/column headers, and grouped formulas. For best results, ensure tables in your documents have clear headers and structure.", icon="ℹ️")
    
    mode_col1, mode_col2 = st.columns([2, 3])
    with mode_col1:
        upload_mode = st.radio(
            "Upload Mode",
            options=["Single Product", "Multiple Variants (compare)"],
            index=0 if st.session_state.mode == 'single' else 1,
            horizontal=True,
            help="Use 'Multiple Variants' to compare formulas across different product versions (e.g., old vs new)"
        )
        new_mode = 'single' if upload_mode == "Single Product" else 'multi'
        if new_mode != st.session_state.mode:
            st.session_state.mode = new_mode
            st.session_state.extraction_result = None
            st.session_state.formulas = []
            st.session_state.variant_results = {}
            st.rerun()

    # ========== SINGLE MODE ==========
    if st.session_state.mode == 'single':
        st.markdown("Upload one or more documents for a **single product**. All files will be combined for extraction.")

        file_uploader_key = f"file_uploader_single_{hash(str(sorted(st.session_state.selected_output_variables)))}"
        uploaded_files = st.file_uploader(
            "Select document(s)",
            type=list(ALLOWED_EXTENSIONS),
            accept_multiple_files=True,
            help=f"Accepts: {', '.join(ALLOWED_EXTENSIONS)}. Max {MAX_FILE_SIZE / (1024*1024):.0f} MB each.",
            key=file_uploader_key
        )

        if uploaded_files:
            for uf in uploaded_files:
                if uf.size > MAX_FILE_SIZE:
                    st.error(f"'{uf.name}' exceeds size limit.")
                else:
                    st.info(f"📄 `{uf.name}` ({uf.size / 1024:.1f} KB)")

            if st.button("Analyze Document(s)", type="primary", key="analyze_single"):
                if not st.session_state.selected_output_variables:
                    st.warning("Please select at least one target formula.")
                else:
                    with st.spinner("Extracting formulas..."):
                        formulas = run_extraction_for_variant(
                            "default",
                            uploaded_files,
                            st.session_state.selected_output_variables,
                            st.session_state.custom_input_variables
                        )
                    if formulas is not None:
                        st.session_state.formulas = formulas
                        st.session_state.extraction_result = True
                        st.session_state.formulas_saved = False
                        st.session_state.editing_formula = -1
                        st.success(f"✅ Extracted {len(formulas)} formulas!")
                        st.rerun()

    # ========== MULTI-VARIANT MODE ==========
    else:
        st.markdown("Upload documents for each product variant. Formulas will be extracted separately and **differences highlighted**.")

        num_col, _ = st.columns([1, 3])
        with num_col:
            num_variants = st.number_input(
                "Number of variants",
                min_value=2,
                max_value=MAX_VARIANTS,
                value=st.session_state.num_variants,
                step=1,
                key="num_variants_input"
            )

        if num_variants != st.session_state.num_variants:
            st.session_state.num_variants = num_variants
            # Extend or trim variant names
            while len(st.session_state.variant_names) < num_variants:
                st.session_state.variant_names.append(f"Variant {len(st.session_state.variant_names) + 1}")
            st.session_state.variant_names = st.session_state.variant_names[:num_variants]
            st.rerun()

        # Ensure variant_names list is properly initialized with enough elements
        while len(st.session_state.variant_names) < num_variants:
            st.session_state.variant_names.append(f"Variant {len(st.session_state.variant_names) + 1}")
        
        variant_files = {}
        all_filled = True

        # Render variant upload panels in a grid (up to 3 per row)
        cols_per_row = min(num_variants, 3)
        rows = (num_variants + cols_per_row - 1) // cols_per_row

        for row_idx in range(rows):
            cols = st.columns(cols_per_row)
            for col_idx in range(cols_per_row):
                variant_idx = row_idx * cols_per_row + col_idx
                if variant_idx >= num_variants:
                    break
                with cols[col_idx]:
                    # Safe access with bounds checking
                    if variant_idx < len(st.session_state.variant_names):
                        default_name = st.session_state.variant_names[variant_idx]
                    else:
                        default_name = f"Variant {variant_idx + 1}"
                    variant_name = st.text_input(
                        f"Variant {variant_idx + 1} Name",
                        value=default_name,
                        key=f"variant_name_{variant_idx}",
                        placeholder="e.g. New Product, Old Product"
                    )
                    # Ensure the list is long enough before assigning
                    while len(st.session_state.variant_names) <= variant_idx:
                        st.session_state.variant_names.append(f"Variant {len(st.session_state.variant_names) + 1}")
                    st.session_state.variant_names[variant_idx] = variant_name

                    uf_key = f"variant_files_{variant_idx}_{hash(str(sorted(st.session_state.selected_output_variables)))}"
                    files = st.file_uploader(
                        f"Upload docs for '{variant_name}'",
                        type=list(ALLOWED_EXTENSIONS),
                        accept_multiple_files=True,
                        key=uf_key,
                        help="You can upload multiple files; they'll be merged for extraction"
                    )
                    variant_files[variant_name] = files or []
                    if not files:
                        all_filled = False
                    else:
                        total_size = sum(f.size for f in files)
                        st.caption(f"📎 {len(files)} file(s), {total_size / 1024:.1f} KB total")

        if not all_filled:
            st.info("⬆️ Please upload at least one document for each variant to enable extraction.")

        if st.button("Analyze All Variants", type="primary", key="analyze_multi", disabled=not all_filled):
            if not st.session_state.selected_output_variables:
                st.warning("Please select at least one target formula.")
            else:
                st.session_state.variant_results = {}
                for v_name in st.session_state.variant_names[:num_variants]:
                    files = variant_files.get(v_name, [])
                    if files:
                        st.markdown(f"#### 🔍 Extracting formulas for **{v_name}**...")
                        formulas = run_extraction_for_variant(
                            v_name,
                            files,
                            st.session_state.selected_output_variables,
                            st.session_state.custom_input_variables
                        )
                        if formulas is not None:
                            st.session_state.variant_results[v_name] = formulas
                            st.success(f"✅ {v_name}: {len(formulas)} formulas extracted")
                st.session_state.extraction_result = True
                st.rerun()

    st.markdown("---")

    # ========== RESULTS: SINGLE MODE ==========
    if st.session_state.mode == 'single' and st.session_state.extraction_result and st.session_state.formulas:
        formulas = st.session_state.formulas

        st.markdown(
            f"""
            <div style="
                background: linear-gradient(135deg, #6ec6ff 0%, #005eac 100%);
                padding: 12px 20px;
                border-radius: 8px;
                margin: 16px 0;
            ">
                <h3 style="color: white; margin: 0; font-size: 1.3rem;">
                    📋 Extracted Formulas ({len(formulas)})
                </h3>
            </div>
            """,
            unsafe_allow_html=True
        )

        
        # Compact stats in single line
        col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4)
        with col_stat1:
            st.metric("Total", len(formulas))
        with col_stat2:
            conditional_count = sum(1 for f in formulas if f.get('is_conditional', False))
            st.metric("Conditional", conditional_count)
        with col_stat3:
            with_evidence = sum(1 for f in formulas if f.get('document_evidence', '') not in ('', 'INFERRED', 'No evidence found', 'No supporting evidence found'))
            st.metric("With Evidence", with_evidence)
        with col_stat4:
            st.metric("Target Variables", len(st.session_state.selected_output_variables))

        tab_view, tab_edit = st.tabs(["👁️ View", "✏️ Edit"])

        with tab_view:
            for formula in formulas:
                render_formula_card(formula)

        with tab_edit:
            col_header1, col_header2, col_header3 = st.columns([3, 5, 2])
            with col_header1:
                st.markdown("**Formula Name**")
            with col_header2:
                st.markdown("**Expression**")
            with col_header3:
                st.markdown("**Actions**")

            for i, formula in enumerate(st.session_state.formulas):
                col1, col2, col3 = st.columns([3, 5, 2])
                with col1:
                    if st.session_state.editing_formula == i:
                        new_name = st.text_input("Formula Name", value=formula.get("formula_name", ""), key=f"name_{i}", label_visibility="collapsed")
                    else:
                        st.markdown(f"**{formula.get('formula_name', '')}**")
                with col2:
                    if st.session_state.editing_formula == i:
                        new_expression = st.text_area("Expression", value=formula.get("formula_expression", ""), key=f"expr_{i}", label_visibility="collapsed", height=68)
                    else:
                        st.code(formula.get("formula_expression", ""), language="python")
                with col3:
                    if st.session_state.editing_formula == i:
                        c1, c2 = st.columns(2)
                        with c1:
                            if st.button("Save", key=f"save_{i}"):
                                st.session_state.formulas[i]["formula_name"] = new_name
                                st.session_state.formulas[i]["formula_expression"] = new_expression
                                st.session_state.formulas_saved = False
                                st.session_state.editing_formula = -1
                                st.rerun()
                        with c2:
                            if st.button("Cancel", key=f"cancel_{i}"):
                                st.session_state.editing_formula = -1
                                st.rerun()
                    else:
                        c1, c2 = st.columns(2)
                        with c1:
                            if st.button("Edit", key=f"edit_{i}"):
                                st.session_state.editing_formula = i
                                st.rerun()
                        with c2:
                            if st.button("Delete", key=f"delete_{i}"):
                                st.session_state.formulas.pop(i)
                                st.session_state.editing_formula = -1
                                st.rerun()

                if i < len(st.session_state.formulas) - 1:
                    st.markdown('<hr style="margin: 0.5rem 0; border: 0; border-top: 1px solid #e0e0e0;">', unsafe_allow_html=True)

            st.markdown("---")
            st.markdown("#### Add New Formula")
            ca1, ca2, ca3 = st.columns([3, 5, 2])
            with ca1:
                new_formula_name = st.text_input("New Formula Name", key="new_formula_name", label_visibility="collapsed", placeholder="Formula name")
            with ca2:
                new_formula_expression = st.text_area("New Formula Expression", key="new_formula_expression", label_visibility="collapsed", placeholder="Formula expression", height=68)
            with ca3:
                st.write("")
                if st.button("Add Formula", key="add_new_formula"):
                    if new_formula_name.strip() and new_formula_expression.strip():
                        st.session_state.formulas.append({"formula_name": new_formula_name.strip(), "formula_expression": new_formula_expression.strip()})
                        st.session_state.formulas_saved = False
                        st.rerun()

        st.markdown("---")
        st.markdown("""<div style="background: #e8f4f8; padding: 10px 16px; border-radius: 6px; border-left: 3px solid #2196f3;">
            <strong style="color: #1565c0;">⬇️ Export Formulas</strong>
        </div>""", unsafe_allow_html=True)
        st.markdown('<div style="padding: 8px 0;"></div>', unsafe_allow_html=True)
        
        export_data = {"total_formulas": len(st.session_state.formulas), "formulas": st.session_state.formulas}
        col_exp1, col_exp2 = st.columns(2)
        with col_exp1:
            st.download_button("📥 JSON (Full)", data=json.dumps(export_data, indent=2, default=str), file_name="extracted_formulas.json", mime="application/json", use_container_width=True)
        with col_exp2:
            # Comprehensive CSV export with all fields
            csv_rows = []
            for f in st.session_state.formulas:
                csv_rows.append({
                    "Formula Name": f.get("formula_name", ""),
                    "Expression": f.get("formula_expression", ""),
                    "Is Conditional": f.get("is_conditional", False),
                    "Conditions": json.dumps(f.get("conditions", []), default=str) if f.get("conditions") else "",
                    "Business Context": f.get("business_context", ""),
                    "Document Evidence": f.get("document_evidence", ""),
                    "Source Method": f.get("source_method", ""),
                    "Variants Info": f.get("variants_info", ""),
                    "Specific Variables": json.dumps(f.get("specific_variables", {}), default=str) if f.get("specific_variables") else ""
                })
            csv_data = pd.DataFrame(csv_rows).to_csv(index=False)
            st.download_button("📥 CSV (Full)", data=csv_data, file_name="extracted_formulas.csv", mime="text/csv", use_container_width=True)

    # ========== RESULTS: MULTI-VARIANT MODE ==========
    elif st.session_state.mode == 'multi' and st.session_state.extraction_result and st.session_state.variant_results:
        variant_results = st.session_state.variant_results
        variant_names = [n for n in st.session_state.variant_names[:st.session_state.num_variants] if n in variant_results]

        if not variant_names:
            st.warning("No variant results available.")
        else:
            # Compact header
            st.markdown(f"""<div style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); padding: 12px 20px; border-radius: 8px; margin: 16px 0;">
                <h3 style="color: white; margin: 0; font-size: 1.3rem;">📊 Multi-Variant Comparison ({len(variant_names)} variants)</h3>
            </div>""", unsafe_allow_html=True)

            # --- COMPARISON TABLE ---
            with st.expander("🔍 Side-by-Side Comparison", expanded=True):
                df = build_formula_comparison_table(variant_results, variant_names)

                def highlight_diffs(row):
                    styles = [''] * len(row)
                    if row.get('_has_diff', False):
                        styles = ['background-color: #fde8e8; color: #c0392b; font-weight: bold'] * len(row)
                    return styles

                display_df = df.drop(columns=['_has_diff'])
                styled = display_df.style.apply(
                    lambda row: ['background-color: #fde8e8; color: #c0392b; font-weight: bold' if df.loc[row.name, '_has_diff'] else '' for _ in row],
                    axis=1
                )
                st.dataframe(styled, use_container_width=True, hide_index=True)

                # Compact Legend
                st.caption("🔴 Red = Differs across variants | ✅ Normal = Same across all variants")

            # --- PER-VARIANT DETAILED VIEW ---
            st.markdown("**📄 Detailed Formulas per Variant**")

            # Figure out which formula names differ
            all_formula_names = []
            for vn in variant_names:
                for f in variant_results[vn]:
                    fn = f.get("formula_name", "")
                    if fn and fn not in all_formula_names:
                        all_formula_names.append(fn)

            differing_formulas = set()
            for fn in all_formula_names:
                exprs = []
                for vn in variant_names:
                    match = next((f for f in variant_results[vn] if f.get("formula_name") == fn), None)
                    exprs.append(match["formula_expression"] if match else "—")
                unique = set(e for e in exprs if e != "—")
                if len(unique) > 1:
                    differing_formulas.add(fn)

            tabs = st.tabs(variant_names)
            for tab, vn in zip(tabs, variant_names):
                with tab:
                    formulas = variant_results[vn]
                    n_diff = sum(1 for f in formulas if f.get("formula_name") in differing_formulas)
                    st.caption(f"{len(formulas)} formulas extracted — {n_diff} differ from at least one other variant")

                    # Filter toggle
                    show_only_diffs = st.checkbox(f"Show only differing formulas", key=f"diff_filter_{vn}")

                    for formula in formulas:
                        fn = formula.get("formula_name", "")
                        has_diff = fn in differing_formulas
                        if show_only_diffs and not has_diff:
                            continue
                        render_formula_card(formula, variant_name=vn, highlight=has_diff)

            # --- DIFF SUMMARY ---
            st.markdown("---")
            st.markdown("""<div style="background: #fff3cd; padding: 10px 16px; border-radius: 6px; border-left: 3px solid #ffc107;">
                <strong style="color: #856404;">🔴 Differences Summary</strong>
            </div>""", unsafe_allow_html=True)
            st.markdown('<div style="padding: 4px 0;"></div>', unsafe_allow_html=True)
            
            if differing_formulas:
                st.markdown(f"**{len(differing_formulas)}** formula(s) differ across variants:")
                for fn in sorted(differing_formulas):
                    with st.expander(f"⚠️ {fn}", expanded=False):
                        for vn in variant_names:
                            match = next((f for f in variant_results[vn] if f.get("formula_name") == fn), None)
                            st.markdown(f"**{vn}:**")
                            if match:
                                st.code(match.get("formula_expression", "—"), language="python")
                                ev = match.get("document_evidence", "")
                                if ev and ev not in ("No supporting evidence found", "No evidence found", "INFERRED"):
                                    st.caption(f"📄 Evidence: {ev[:300]}...")
                            else:
                                st.markdown("_Not found in this variant_")
                        st.markdown("---")
            else:
                st.success("✅ All formulas are identical across all variants!")

            # --- EXPORT ---
            st.markdown("---")
            st.markdown("""<div style="background: #e8f4f8; padding: 10px 16px; border-radius: 6px; border-left: 3px solid #2196f3;">
                <strong style="color: #1565c0;">⬇️ Export Comparison</strong>
            </div>""", unsafe_allow_html=True)
            st.markdown('<div style="padding: 8px 0;"></div>', unsafe_allow_html=True)
            
            export_all = {
                "variants": variant_names,
                "differing_formulas": list(differing_formulas),
                "results": {vn: variant_results[vn] for vn in variant_names}
            }
            col_e1, col_e2 = st.columns(2)
            with col_e1:
                st.download_button("📥 JSON (All)", data=json.dumps(export_all, indent=2, default=str), file_name="multi_variant_formulas.json", mime="application/json", use_container_width=True)
            with col_e2:
                # Comprehensive CSV with all fields and variant column
                flat_rows = []
                for vn in variant_names:
                    for f in variant_results[vn]:
                        flat_rows.append({
                            "Variant": vn,
                            "Formula Name": f.get("formula_name", ""),
                            "Expression": f.get("formula_expression", ""),
                            "Is Conditional": f.get("is_conditional", False),
                            "Conditions": json.dumps(f.get("conditions", []), default=str) if f.get("conditions") else "",
                            "Business Context": f.get("business_context", ""),
                            "Document Evidence": f.get("document_evidence", ""),
                            "Source Method": f.get("source_method", ""),
                            "Variants Info": f.get("variants_info", ""),
                            "Specific Variables": json.dumps(f.get("specific_variables", {}), default=str) if f.get("specific_variables") else "",
                            "Differs": f.get("formula_name", "") in differing_formulas,
                        })
                st.download_button("📥 CSV (Full)", data=pd.DataFrame(flat_rows).to_csv(index=False), file_name="formula_comparison.csv", mime="text/csv", use_container_width=True)

    # ========== PROCEED TO NEXT STEP ==========
    if (st.session_state.mode == 'single' and st.session_state.extraction_result and st.session_state.formulas) or \
       (st.session_state.mode == 'multi' and st.session_state.extraction_result and st.session_state.variant_results):
        st.markdown("---")

        
        col_proceed1, col_proceed2, col_proceed3 = st.columns([1, 2, 1])
        with col_proceed2:
            if st.button("➡️ Proceed to Variable Mapping", type="primary", use_container_width=True, key="proceed_to_mapping"):
                st.switch_page("pages/2_Variable_Mapping.py")

    st.markdown("---")
    st.markdown(
        """<div style="text-align: center; margin-top: 30px; color: #7f8c8d; font-size: 0.9em;">
            <p>Developed by BLIC GenAI team</p>
        </div>""",
        unsafe_allow_html=True
    )


if __name__ == "__main__":
    main()